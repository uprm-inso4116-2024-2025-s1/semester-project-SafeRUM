import pandas as pd
import requests
from docx import Document

# Load email data from CSV


def load_emails_from_csv(csv_filename):
    df = pd.read_csv(csv_filename, skiprows=2)
    df = df.rename(
        columns={"Unnamed: 3": "Github Username", "Unnamed: 6": "Email Address"}
    )
    email_data = df[["Github Username", "Email Address"]].dropna()
    return dict(zip(email_data["Github Username"], email_data["Email Address"]))


# Get email by GitHub username


def get_user_email_from_csv(username, email_data):
    return email_data.get(username, "No email available")


# Paths and authentication

csv_file_path = "this file (##SafeRUM Members InfoSheet - Hoja 1##) here"
email_data = load_emails_from_csv(csv_file_path)
REPO = "uprm-inso4116-2024-2025-s1/semester-project-SafeRUM"
MILESTONE_NUMBER = 3
TOKEN = "Github token"  # GitHub token

# Initialize document

doc = Document()
doc.add_heading("Log Book - Milestone 3", 0)

# Create table

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Issue Title"
hdr_cells[1].text = "Assignee"
hdr_cells[2].text = "Email"

page = 1
total_issues = 0

while True:
    # Fetch closed issues

    params = {
        "milestone": MILESTONE_NUMBER,
        "state": "closed",
        "per_page": 30,
        "page": page,
    }
    response = requests.get(
        f"https://api.github.com/repos/{REPO}/issues",
        headers={"Authorization": f"token {TOKEN}"},
        params=params,
    )

    if response.status_code != 200:
        print(f"Error {response.status_code}: {response.text}")
        break
    issues = response.json()
    if not issues:
        break
    # Process issues and add assignee info

    for issue in issues:
        if "pull_request" not in issue:
            issue_title = issue["title"]
            assignees = issue["assignees"]

            if not assignees:
                continue
            for assignee in assignees:
                username = assignee["login"]
                email = get_user_email_from_csv(username, email_data)

                row_cells = table.add_row().cells
                row_cells[0].text = issue_title
                row_cells[1].text = username
                row_cells[2].text = email

                print(f"{username}: {email}")
                total_issues += 1
    page += 1
# Save document

output_file = "log_book_milestone_3.docx"
doc.save(output_file)
print(f"Document created with {total_issues} issues: {output_file}")
