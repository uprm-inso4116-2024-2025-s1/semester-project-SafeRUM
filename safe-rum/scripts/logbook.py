import requests
import pandas as pd
from docx import Document

# Load and clean the CSV data
def load_emails_from_csv(csv_filename):
    df = pd.read_csv(csv_filename, skiprows=3)  # Skip the first 3 metadata rows
    email_data = df[['Github Username', 'Email Adress']].dropna()
    return dict(zip(email_data['Github Username'], email_data['Email Adress']))

# Function to get email from the CSV dictionary
def get_user_email_from_csv(username, email_data):
    return email_data.get(username, 'No email available')

# Load email data from the provided CSV file
##### Warning you need the sheet file named "SafeRUM Members InfoSheet" ######
email_data = load_emails_from_csv('Your csv file.csv')

# Configure your repository and milestone number
REPO = "uprm-inso4116-2024-2025-s1/semester-project-SafeRUM"
MILESTONE_NUMBER = 2
TOKEN = "your Github API key"

doc = Document()
doc.add_heading('Log Book - Milestone 2', 0)

# Create a table with headers
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Issue Title'
hdr_cells[1].text = 'Contributor'
hdr_cells[2].text = 'Email'

page = 1
total_issues = 0

while True:
    params = {"milestone": MILESTONE_NUMBER, "state": "closed", "per_page": 30, "page": page}
    response = requests.get(f"https://api.github.com/repos/{REPO}/issues", headers={"Authorization": f"token {TOKEN}"}, params=params)

    if response.status_code != 200:
        print(f"Error {response.status_code}: {response.text}")
        break

    issues = response.json()
    if not issues:
        break

    for issue in issues:
        if "pull_request" not in issue:
            username = issue['user']['login']
            email = get_user_email_from_csv(username, email_data)

            row_cells = table.add_row().cells
            row_cells[0].text = issue['title']
            row_cells[1].text = username
            row_cells[2].text = email

            print(f"{username}: {email}")
            total_issues += 1

    page += 1

doc.save('log_book_milestone_2.docx')
print(f"Document successfully created with {total_issues} issues: log_book_milestone_2.docx")
